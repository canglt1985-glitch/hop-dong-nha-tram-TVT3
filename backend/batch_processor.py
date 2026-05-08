import pandas as pd
import os
from finance_service import FinanceService as fs
from docx import Document
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from num2words import num2words
import sys

# Đảm bảo in được tiếng Việt ra console
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# Cấu hình đường dẫn (macOS - relative to this file's location)
_BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MASTER_EXCEL = os.path.join(_BASE, 'data_source', 'DANH_SACH_HOP_DONG_MAT_BANG_NHA_TRAM_21_04_2026.xlsx')
PRICE_EXCEL  = os.path.join(_BASE, 'data_source', 'MBF ĐNa_BC_VB 1245_BÁO CÁO NGÀY.xlsx')
PHUONG_XA_FILE = os.path.join(_BASE, 'data_source', 'phuong_xa.xlsx')
MASTER_DOC   = os.path.join(_BASE, 'templates', 'MASTER_TEMPLATE_VFINAL.docx')
O_DIR        = os.path.join(_BASE, 'output', 'batch_test')

PHUONG_XA_MAP = {}

def load_phuong_xa():
    global PHUONG_XA_MAP
    if os.path.exists(PHUONG_XA_FILE):
        try:
            df_px = pd.read_excel(PHUONG_XA_FILE)
            for _, row in df_px.iterrows():
                s_id = str(row.iloc[0]).strip()
                val_ao = str(row.iloc[40]).strip() if len(row) > 40 else ""
                if s_id and val_ao and val_ao != "nan":
                    if " - " in val_ao:
                        val_ao = val_ao.split(" - ")[1]
                    PHUONG_XA_MAP[s_id] = val_ao.title()
        except Exception as e:
            print(f"Warning: Could not load phuong xa file: {e}")

# Gọi load ngay khi khởi động
load_phuong_xa()

def apply_font_to_paragraph(p, font_name, font_size):
    for run in p.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)

def safe_float(val):
    try:
        return float(val) if pd.notna(val) else 0.0
    except:
        return 0.0

def get_cycle_start(master_row):
    """Trích xuất ngày bắt đầu từ Cột 29 (Ví dụ 01/01/2026 - 30/06/2026)"""
    val = str(master_row.iloc[29])
    try:
        date_str = val.split("-")[0].strip()
        return datetime.strptime(date_str, "%d/%m/%Y")
    except:
        return datetime(2025, 12, 1)

def get_price_info(site_id, df_price):
    """Lấy thông tin giá từ Price Sheet 1245"""
    res = df_price[df_price.iloc[:, 1].astype(str).str.contains(site_id, na=False)]
    if res.empty: return None
    p_row = res.iloc[0]
    
    mb = safe_float(p_row.iloc[42])
    pm = safe_float(p_row.iloc[43])
    mfd = safe_float(p_row.iloc[50])
    cot = safe_float(p_row.iloc[52])
    giamtru = safe_float(p_row.iloc[62])
    
    new_price_raw = safe_float(p_row.iloc[107])
    if new_price_raw > 0:
        new_price = new_price_raw
    else:
        tong_qd = (mb + pm + mfd + cot) + giamtru
        new_price = (tong_qd // 50000) * 50000
    
    return {
        "MB_QĐ02": mb,
        "PM": pm,
        "MFD": mfd,
        "COT": cot,
        "GIAM_TRU": giamtru,
        "TONG_QD": (mb + pm + mfd + cot) + giamtru,
        "NEW_PRICE": new_price
    }

def get_period_amt(p_start, p_end, old_price, new_price):
    """Tính tiền cho một kỳ, tự động chẻ giá cũ/mới theo mốc 01/10/2025"""
    cutoff = datetime(2025, 10, 1)
    total = 0.0
    
    curr = p_start
    while curr <= p_end:
        next_m = curr + relativedelta(months=1)
        
        if next_m > p_end + relativedelta(days=1):
            days_in_month = (next_m - curr).days
            actual_days = (p_end + relativedelta(days=1) - curr).days
            price = old_price if curr < cutoff else new_price
            total += price * (actual_days / days_in_month)
            break
        else:
            price = old_price if curr < cutoff else new_price
            total += price
            curr = next_m
            
    return round(total, -3)

def get_site_schedule(master_row, df_price):
    """Tính toán lịch trình thanh toán mà không xuất file"""
    site_id = str(master_row.iloc[2]).strip()
    old_price_val = safe_float(master_row.iloc[19])
    
    p_info = get_price_info(site_id, df_price)
    if not p_info: return None
    new_price_val = p_info["NEW_PRICE"]
    
    # Expiry logic
    end_date_raw = master_row.iloc[26]
    if pd.isna(end_date_raw):
        end_contract = datetime(2028, 12, 31)
    else:
        try:
            if isinstance(end_date_raw, datetime): end_contract = end_date_raw
            else: end_contract = datetime.strptime(str(end_date_raw).split(" ")[0], "%Y-%m-%d")
        except: end_contract = datetime(2028, 12, 31)

    # 1. Logic Gia hạn tự động: Nếu hết hạn trước 01/07/2026 -> +5 năm
    threshold_date = datetime(2026, 7, 1)
    original_end_contract = end_contract
    if end_contract < threshold_date:
        end_contract = end_contract + relativedelta(years=5)
    
    # 2. Chu kỳ thanh toán cố định: 06 tháng/lần
    cycle_months = 6

    # Lấy hàng dữ liệu giá trong df_price
    p_row_matches = df_price[df_price.iloc[:, 1].astype(str).str.contains(site_id, na=False)]
    if p_row_matches.empty:
        p_row_matches = df_price[df_price.iloc[:, 0].astype(str).str.contains(site_id, na=False)]
    
    if p_row_matches.empty: return None
    p_row = p_row_matches.iloc[0]
    site_type = str(p_row.iloc[2]).strip() # Cột 2: HTCS
    
    # 1. Lấy TỔNG GIÁ GỐC (Regulated) từ Cột 41
    total_raw_col41 = p_row.iloc[41] if pd.notna(p_row.iloc[41]) else 0
    
    # 2. Quét các thành phần GIÁ THANH TOÁN theo tọa độ AQ-BK
    # Cột 42: Mặt bằng
    mb_raw_aq = p_row.iloc[42] if pd.notna(p_row.iloc[42]) else 0
    
    # Cột 43-48: Phòng máy
    pm_raw_ar_aw = 0
    pm_label_val = "Phòng máy mặt đất"
    for i in range(43, 49):
        val = p_row.iloc[i]
        if pd.notna(val) and isinstance(val, (int, float)) and val > 0:
            pm_raw_ar_aw = val
            title = str(df_price.iloc[5, i]).strip()
            if title and title.lower() != 'nan':
                pm_label_val = title.capitalize()
            break
            
    # Cột 50: Phòng MFĐ
    mfd_raw_ay = p_row.iloc[50] if pd.notna(p_row.iloc[50]) else 0
    
    # Cột 51-52: Cột anten
    # Cột 51-54: Cột anten
    cot_raw_az_ba = 0
    for i in range(51, 54):
        val = p_row.iloc[i]
        if pd.notna(val) and isinstance(val, (int, float)) and val > 0:
            cot_raw_az_ba = val
            break
    
    # Cột 62: Giảm trừ dùng chung
    giam_tru_raw_bk = p_row.iloc[62] if pd.notna(p_row.iloc[62]) else 0
    
    # 3. LẤY CÁC HẠNG MỤC PHÁT SINH KHÁC
    processed_cols = [41, 42, 50, 62] + list(range(43, 49)) + list(range(51, 54))
    other_items = []
    other_items_pay_total = 0
    
    for i in range(43, 62): 
        if i in processed_cols: continue
        val = p_row.iloc[i]
        if pd.notna(val) and isinstance(val, (int, float)) and val > 0:
            raw_val = val
            pay_val = (int(float(raw_val)) // 50000) * 50000
            title = str(df_price.iloc[5, i]).strip()
            if title and title.lower() != 'nan':
                other_items.append({
                    "title": title.capitalize(),
                    "raw": raw_val,
                    "pay": pay_val
                })
                other_items_pay_total += pay_val

    # 4. LOGIC LÀM TRÒN HAI LỚP (Theo yêu cầu: 1256000 còn 1250000)
    # Lớp 1: Làm tròn hạng mục Cột (bao gồm giảm trừ dùng chung)
    cot_price_pay = ( (int(float(cot_raw_az_ba)) + int(float(giam_tru_raw_bk))) // 50000 ) * 50000
    
    # Lớp 2: Tính tổng tạm thời và làm tròn xuống cho TỔNG CHỐT
    temp_total = mb_raw_aq + pm_raw_ar_aw + mfd_raw_ay + cot_price_pay + sum(item["raw"] for item in other_items)
    new_price_val = (int(temp_total) // 50000) * 50000
    
    # Lớp 3: Làm tròn phòng máy và dồn phần còn lại vào Mặt bằng
    pm_price_pay = (int(float(pm_raw_ar_aw)) // 50000) * 50000 if pm_raw_ar_aw > 0 else 0
    mfd_price_pay = (int(float(mfd_raw_ay)) // 50000) * 50000 if mfd_raw_ay > 0 else 0
    mb_price_pay = new_price_val - pm_price_pay - mfd_price_pay - cot_price_pay - other_items_pay_total

    p_info = {
        "MB_QĐ02": mb_price_pay,
        "PM": pm_price_pay,
        "PM_LABEL": pm_label_val,
        "MFD": mfd_price_pay,
        "COT": cot_price_pay,
        "GIAM_TRU": 0, # Đã gộp vào COT để làm tròn lớp 1
        "TONG_QD": total_raw_col41,
        "NEW_PRICE": new_price_val,
        "other_items": other_items,
        "_RAW_": {
            "MB_QĐ02": mb_raw_aq,
            "PM": pm_raw_ar_aw,
            "MFD": mfd_raw_ay,
            "COT": cot_raw_az_ba,
            "GIAM_TRU": giam_tru_raw_bk,
            "TONG_QD": total_raw_col41
        }
    }

    # LOGIC NGÀY BẮT ĐẦU: Không fix 01/10/2025 mà nhảy theo ngày đã thanh toán + 1
    paid_until_date = master_row.iloc[28] # Cột AC
    if pd.isna(paid_until_date):
        paid_until_date = datetime(2025, 12, 31)
    
    # Ngày bắt đầu tính kỳ tiếp theo = Ngày đã trả xong + 1 ngày
    # (Nhưng mốc 01/10/2025 vẫn dùng để tính khấu trừ nếu có)
    fixed_calc_start = datetime(2025, 10, 1)
    effective_date = paid_until_date + relativedelta(days=1)
    curr_start = effective_date
    
    # Tính khấu trừ (Vẫn tính từ 01/10/2025 đến ngày đã thanh toán)
    deduction_val = 0
    if paid_until_date >= fixed_calc_start:
        rel = relativedelta(paid_until_date + relativedelta(days=1), fixed_calc_start)
        total_months_diff = rel.years * 12 + rel.months + (rel.days / 30.44)
        deduction_val = round((old_price_val - new_price_val) * total_months_diff)

    # TÍNH TOÁN CÁC KỲ (START TỪ NGÀY TIẾP THEO SAU KHI HẾT NỢ)
    periods = []
    total_amount = 0
    c_no = 1
    
    # Logic mới: Luôn chạy tròn kỳ 6 tháng, không cắt ngắn kỳ cuối
    while curr_start < end_contract:
        # Mỗi kỳ luôn là tròn 6 tháng (ép theo yêu cầu: chu kỳ cuối cũng tròn 6 tháng)
        curr_end = curr_start + relativedelta(months=cycle_months) - relativedelta(days=1)
        
        amt = get_period_amt(curr_start, curr_end, old_price_val, new_price_val)
        if c_no == 1 and deduction_val > 0:
            amt -= deduction_val
            
        total_amount += amt
        periods.append({
            "no": c_no,
            "start": curr_start,
            "end": curr_end,
            "amount": amt
        })
        
        # Cập nhật ngày bắt đầu kỳ tiếp theo
        curr_start = curr_end + relativedelta(days=1)
        c_no += 1
    
    # Cập nhật lại ngày kết thúc hợp đồng thực tế bằng ngày cuối của kỳ cuối cùng
    if periods:
        end_contract = periods[-1]["end"]
        
    return {
        "site_id": site_id,
        "old_price": old_price_val,
        "new_price": new_price_val,
        "end_contract": end_contract,
        "original_end_contract": original_end_contract,
        "periods": periods,
        "total_amount": total_amount,
        "deduction": deduction_val,
        "paid_until": paid_until_date,
        "p_info": p_info
    }

def process_site(master_row, df_price):
    site_id = str(master_row.iloc[2]).strip()
    contract_no = str(master_row.iloc[1]).strip()
    owner_name = str(master_row.iloc[10]).strip()
    
    # Sử dụng logic tính toán chung
    schedule = get_site_schedule(master_row, df_price)
    if not schedule: return False, "Not found in Price Sheet"
    
    new_price_val = schedule["new_price"]
    old_price_val = schedule["old_price"]
    p_info = schedule["p_info"]
    deduction_val = schedule["deduction"]
    paid_until_date = schedule["paid_until"]
    periods = schedule["periods"]
    total_amount = schedule["total_amount"]
    
    # Khai báo các loại phí từ p_info
    mb_price = p_info["MB_QĐ02"]
    bh_pm_price = p_info["PM"]
    mfd_price = p_info["MFD"]
    cot_price = p_info["COT"]

    # 1. Đếm chu kỳ thanh toán còn lại (từ hôm nay)
    today = datetime.now()
    remaining_cycles = len([p for p in periods if p["end"] >= today])
    
    # 2. Kiểm tra khớp chủ tài khoản (So sánh OWNER_NAME và ACCOUNT_OWNER thực tế)
    # Lấy tên chủ tài khoản thực tế từ Cột 32 (Master)
    bank_owner_name = str(master_row.iloc[32]).strip()
    is_owner_match = owner_name.lower() in bank_owner_name.lower() or bank_owner_name.lower() in owner_name.lower()
    owner_match_status = "✅ Khớp" if is_owner_match else f"⚠️ Lệch (Bank: {bank_owner_name})"

    # Số tài khoản lấy từ Cột 33
    bank_acc_no = str(master_row.iloc[33]).strip()

    # 3. Kiểm tra gia hạn hợp đồng bằng ngày nguyên thủy
    end_contract = schedule["end_contract"]
    original_end_contract = schedule.get("original_end_contract", end_contract)
    needs_extension = original_end_contract < datetime(2026, 7, 1)
    extension_status = "🔴 Cần gia hạn" if needs_extension else "🟢 Còn hạn"

    # XỬ LÝ XÃ MỚI
    suggested_new = PHUONG_XA_MAP.get(site_id, "")
    address_new_val = ""
    if suggested_new:
        address_new_val = suggested_new
        if "Đồng Nai" not in address_new_val:
            address_new_val += ", Đồng Nai"

    paid_until_str = paid_until_date.strftime("%d/%m/%Y")
    if deduction_val > 0:
        deduction_text = f"Giảm trừ số tiền {fs.format_vn_currency(deduction_val)} đồng đã thanh toán từ 01/10/2025 đến {paid_until_str} vào kỳ thanh toán tiếp theo."
    else:
        deduction_text = ""

    replacements = {
        "{{SITE_ID}}": site_id,
        "{{CONTRACT_NO}}": contract_no,
        "{{CONTRACT_DATE}}": master_row.iloc[25].strftime('%d/%m/%Y') if pd.notna(master_row.iloc[25]) and isinstance(master_row.iloc[25], datetime) else str(master_row.iloc[25]),
        "{{OWNER_NAME}}": owner_name,
        "{{ADDRESS_OLD}}": str(master_row.iloc[7]).strip() if pd.notna(master_row.iloc[7]) else "",
        " ({{ADDRESS_NEW}})": f" ({address_new_val})" if address_new_val else "", 
        "{{ADDRESS_NEW}}": address_new_val, 
        "{{CONTACT_ADDR}}": str(master_row.iloc[8]).strip() if pd.notna(master_row.iloc[8]) else "",
        "{{PHONE}}": str(master_row.iloc[12]).strip() if pd.notna(master_row.iloc[12]) else "",
        "{{ACCOUNT_NO}}": bank_acc_no,
        "{{ACCOUNT_OWNER}}": str(master_row.iloc[32]).strip() if pd.notna(master_row.iloc[32]) else owner_name,
        "{{BANK_NAME}}": f"{master_row.iloc[34]} - {master_row.iloc[35]}" if pd.notna(master_row.iloc[34]) else "",
        "{{OLD_PRICE}}": fs.format_vn_currency(old_price_val),
        "{{NEW_PRICE}}": fs.format_vn_currency(new_price_val),
        "{{P_MB}}": fs.format_vn_currency(p_info["MB_QĐ02"]),
        "{{P_PM}}": fs.format_vn_currency(p_info["PM"]),
        "{{P_MFD}}": fs.format_vn_currency(p_info["MFD"]),
        "{{P_COT}}": fs.format_vn_currency(p_info["COT"]),
        "{{NEW_PRICE_TEXT}}": fs.format_to_text(new_price_val),
        "{{REPORT_DATE}}": today.strftime('%d/%m/%Y'),
        "{{DEDUCTION_TEXT}}": deduction_text,
        "{{MB_QĐ02}}": fs.format_vn_currency(p_info["_RAW_"]["MB_QĐ02"]),
        "{{MFĐ_1245}}": fs.format_vn_currency(p_info["_RAW_"]["MFD"]),
        "{{PM_1245}}": fs.format_vn_currency(p_info["_RAW_"]["PM"]),
        "{{COT_1245}}": fs.format_vn_currency(p_info["_RAW_"]["COT"]),
        "{{GIAM_TRU}}": fs.format_vn_currency(p_info["_RAW_"]["GIAM_TRU"]),
        "{{COT_CHOT}}": fs.format_vn_currency(p_info["COT"]),
        "{{TONG_QD}}": fs.format_vn_currency(p_info["_RAW_"]["TONG_QD"]),
        "{{TONG_CHOT}}": fs.format_vn_currency(new_price_val),
        "Phòng máy mặt đất": p_info["PM_LABEL"],
        "{{PM_LABEL}}": p_info["PM_LABEL"],
        "{{TL_COT}}": f"{((p_info['COT'] / (p_info['_RAW_']['COT'] + p_info['_RAW_']['GIAM_TRU'])) - 1) * 100:.2f}%" if (p_info['_RAW_']['COT'] + p_info['_RAW_']['GIAM_TRU']) != 0 else "0%", 
        "{{TL_PM}}": f"{((p_info['PM'] / p_info['_RAW_']['PM']) - 1) * 100:.2f}%" if p_info['_RAW_']['PM'] != 0 else "0%",
        "{{TL_MB}}": f"{((p_info['MB_QĐ02'] / p_info['_RAW_']['MB_QĐ02']) - 1) * 100:.2f}%" if p_info['_RAW_']['MB_QĐ02'] != 0 else "0%",
        "{{TL_MFD}}": f"{((p_info['MFD'] / p_info['_RAW_']['MFD']) - 1) * 100:.2f}%" if p_info['_RAW_']['MFD'] != 0 else "0%",
        "{{TL_TONG}}": f"{((new_price_val / p_info['_RAW_']['TONG_QD']) - 1) * 100:.2f}%" if p_info['_RAW_']['TONG_QD'] != 0 else "0%",
        "{{END_DATE}}": end_contract.strftime('%d/%m/%Y'),
        "{{TOTAL_AMOUNT}}": fs.format_vn_currency(total_amount),
        "{{CYCLES_REMAINING}}": str(remaining_cycles),
        "{{OWNER_MATCH}}": owner_match_status,
        "{{EXTENSION_STATUS}}": extension_status,
        "{{PAY_ROW}}": ""
    }

    doc = Document(MASTER_DOC)
    
    # Cập nhật các hạng mục phát sinh vào bảng giá (Table 1)
    if schedule["p_info"].get("other_items"):
        price_table = doc.tables[1]
        for item in schedule["p_info"]["other_items"]:
            new_row = price_table.add_row()
            price_table._tbl.insert(-2, new_row._tr)
            
            new_row.cells[0].text = item["title"]
            new_row.cells[1].text = "" 
            new_row.cells[2].text = fs.format_vn_currency(item["raw"])
            new_row.cells[3].text = "" 
            new_row.cells[4].text = fs.format_vn_currency(item["pay"])
            ratio_val = f"{((item['pay'] / item['raw']) - 1) * 100:.2f}%" if item['raw'] != 0 else "0%"
            new_row.cells[5].text = ratio_val
            
            for cell in new_row.cells:
                for p in cell.paragraphs:
                    apply_font_to_paragraph(p, 'Times New Roman', 14)

    import re
    total_replacement = f"Tổng cộng:             {fs.format_vn_currency(total_amount)}"
    
    for p in doc.paragraphs:
        # Tìm các dòng có "Tổng cộng" và con số phía sau (kể cả 142.800.000 hay bất kỳ số nào)
        if "Tổng cộng:" in p.text:
            # Nếu dòng này ở NGOÀI bảng và có chứa số, ta thay bằng tổng mới
            if re.search(r'[\d\.]+', p.text):
                p.text = total_replacement
                apply_font_to_paragraph(p, 'Times New Roman', 14)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def replace_in_paragraph(p):
        text = p.text
        original_text = text
        for old, new in replacements.items():
            if old in text:
                text = text.replace(old, str(new))
        
        # Thay thế ngày 30/09/2028 cứng (nếu có) bằng ngày kết thúc thực tế
        if "30/09/2028" in text:
            text = text.replace("30/09/2028", end_contract.strftime('%d/%m/%Y'))
        
        if text != original_text:
            if "Một bên là: Ông/Bà " in text and "(Gọi tắt là bên A)" in text:
                p.clear()
                r1 = p.add_run(f"Một bên là: Ông/Bà ")
                r2 = p.add_run(owner_name)
                r2.bold = True
                r3 = p.add_run(" (Gọi tắt là bên A)")
            else:
                p.text = text
            apply_font_to_paragraph(p, 'Times New Roman', 14)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    replace_in_paragraph(p)

    pay_table = doc.tables[4]
    
    for i, period in enumerate(periods):
        c_no = period["no"]
        start_date = period["start"]
        c_end = period["end"]
        amt = period["amount"]
        
        line = f"+ Kỳ {c_no}: từ ngày {start_date.strftime('%d/%m/%Y')}  đến ngày {c_end.strftime('%d/%m/%Y')}  Số tiền là: {fs.format_vn_currency(amt)}"
        print(f"   {line}")
        
        if i == 0:
            row = pay_table.rows[0]
        else:
            row = pay_table.add_row()
            
        cell = row.cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(line)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(14)

    # Đã loại bỏ đoạn code thừa tự động thêm một hàng Tổng cộng vào cuối bảng thanh toán,
    # vì Template Word đã có sẵn hoặc có một đoạn paragraph thay thế bên ngoài bảng.

    suffix = "_CAN_GIA_HAN" if needs_extension else ""
    out_file = os.path.join(O_DIR, f"{site_id}_PHU_LUC{suffix}.docx")
    doc.save(out_file)
    return True, out_file

def run_batch(site_filter=None, limit=None):
    df_price = pd.read_excel(PRICE_EXCEL, sheet_name='Chi tiết thuê trạm', header=None)
    df_master = pd.read_excel(MASTER_EXCEL)
    
    if site_filter:
        df_master = df_master[df_master.iloc[:, 2].astype(str).str.contains(site_filter, na=False, case=False)]
    
    count = 0
    exported_count = 0
    for idx, row in df_master.iterrows():
        site_id = str(row.iloc[2]).strip()
        if site_id == "nan": continue
        
        print(f"\n[{exported_count + 1}] Processing: {site_id}")
        ok, res = process_site(row, df_price)
        if ok:
            exported_count += 1
            print(f"   [OK] Done: {res}")
        else:
            print(f"   [ERROR]: {res}")
            
        if limit and exported_count >= limit:
            break

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--site", type=str)
    parser.add_argument("--limit", type=int)
    args = parser.parse_args()
    run_batch(site_filter=args.site, limit=args.limit)
