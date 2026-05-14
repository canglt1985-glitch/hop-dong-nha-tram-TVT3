import os
import docx
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx.oxml.ns import qn
import copy
import calendar
from utils.pricing_utils import calculate_aligned_cycles

def format_vn_currency(amount):
    try:
        if amount is None or amount == "": return "0"
        return "{:,.0f}".format(float(amount)).replace(",", ".")
    except Exception:
        return "0"

def number_to_vietnamese_words(number):
    try: number = int(number)
    except: return ""
    if number == 0: return "Không"
    units = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    def read_three_digits(n, is_first=False):
        hundreds = n // 100
        tens = (n % 100) // 10
        ones = n % 10
        res = []
        if hundreds > 0: res.append(units[hundreds] + " trăm")
        elif not is_first: res.append("không trăm")
        if tens > 0:
            if tens == 1: res.append("mười")
            else: res.append(units[tens] + " mươi")
        elif (hundreds > 0 or not is_first) and ones > 0:
            res.append("lẻ")
        if ones > 0:
            if ones == 1 and tens > 1: res.append("mốt")
            elif ones == 5 and tens > 0: res.append("lăm")
            else: res.append(units[ones])
        return " ".join(res)
    groups = []
    temp = number
    while temp > 0:
        groups.append(temp % 1000)
        temp //= 1000
    group_units = ["", "nghìn", "triệu", "tỷ", "nghìn tỷ", "triệu tỷ"]
    words = []
    for idx, val in enumerate(groups):
        if val > 0:
            part = read_three_digits(val, idx == len(groups) - 1)
            unit = group_units[idx]
            if unit: words.append(part + " " + unit)
            else: words.append(part)
    words.reverse()
    text = " ".join(words)
    text = " ".join(text.split()).strip()
    return text[0].upper() + text[1:]

def replace_tag_in_runs(p, tag, replacement):
    if tag not in p.text: return False
    replaced = False
    for run in p.runs:
        if tag in run.text:
            run.text = run.text.replace(tag, str(replacement))
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
            replaced = True
    if not replaced:
        p.text = p.text.replace(tag, str(replacement))
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
        replaced = True
    return replaced

def calculate_contract_dates(expiry_date_val):
    try:
        if isinstance(expiry_date_val, datetime): 
            exp_date = expiry_date_val
        else:
            date_str = str(expiry_date_val).strip()
            try:
                # Thử format DD/MM/YYYY
                exp_date = datetime.strptime(date_str[:10], "%d/%m/%Y")
            except ValueError:
                # Thử format YYYY-MM-DD
                exp_date = datetime.strptime(date_str[:10], "%Y-%m-%d")
                
        # Ngày bắt đầu HĐ mới = Ngày sau ngày hết hạn cũ
        start_date = exp_date + relativedelta(days=1)
        
        # Ngày kết thúc chuẩn hóa:
        is_exact_month_start = start_date.day == 1
        target_end_year = start_date.year + 5
        
        if is_exact_month_start:
            # 01/06/2026 -> 31/05/2031
            target_end_month = start_date.month - 1
            if target_end_month == 0:
                target_end_month = 12
                target_end_year -= 1
        else:
            # 09/04/2026 -> 30/04/2031
            target_end_month = start_date.month
            
        last_day = calendar.monthrange(target_end_year, target_end_month)[1]
        end_date = datetime(target_end_year, target_end_month, last_day)
        
        return start_date, end_date
    except Exception:
        return datetime(2026, 4, 9), datetime(2031, 4, 30)

def generate_document_from_cloud(site_data, template_path, output_dir, prefix, template_type=None):
    site_id = site_data.get("site_id", "")
    owner_name = site_data.get("owner", "")
    
    # Coordinates & Address
    latitude = str(site_data.get("latitude", ""))
    longitude = str(site_data.get("longitude", ""))
    address_old = site_data.get("address_old", "")
    
    banking = site_data.get("banking_info", {})
    account_owner = banking.get("account_owner", "") or owner_name
    account_no = banking.get("account_no", "")
    bank_name = banking.get("bank_name", "")
    bank_branch = banking.get("bank_branch", "")
    full_bank_name = f"{bank_name} - {bank_branch}" if bank_branch else bank_name

    # Calculate dates
    raw_expiry = site_data.get("end_date", "")
    start_date_dt, extended_end_date_dt = calculate_contract_dates(raw_expiry)
    
    # Extract exact original end_date
    original_end_date_dt = None
    if raw_expiry:
        try:
            if isinstance(raw_expiry, datetime): original_end_date_dt = raw_expiry
            else: original_end_date_dt = datetime.strptime(str(raw_expiry)[:10], "%d/%m/%Y")
        except Exception:
            try:
                original_end_date_dt = datetime.strptime(str(raw_expiry)[:10], "%Y-%m-%d")
            except Exception:
                pass
                
    start_date_str = start_date_dt.strftime("%d/%m/%Y")
    end_date_str = extended_end_date_dt.strftime("%d/%m/%Y")
    
    # Always use the extended 5-year date for the contract term in these templates
    end_contract_dt = extended_end_date_dt
    end_contract = end_contract_dt
    
    # OLD_END_DATE calculation
    if original_end_date_dt:
        old_end_date_str = original_end_date_dt.strftime("%d/%m/%Y")
    else:
        old_end_date_str = str(raw_expiry) if raw_expiry else ""
    
    # Prices
    prices = site_data.get("prices_breakdown", {})
    mb_raw = float(prices.get("mb", 0))
    pm_raw = float(prices.get("pm", 0))
    mfd_raw = float(prices.get("mfd", 0))
    cot_raw = float(prices.get("cot", 0))
    giam_tru_raw = float(prices.get("giam_tru", 0))
    vhkt_raw = float(prices.get("vhkt", 0))
    
    # 4. LOGIC LÀM TRÒN HAI LỚP (Theo yêu cầu: 1256000 còn 1250000)
    # Lớp 1: Làm tròn hạng mục Cột (bao gồm giảm trừ dùng chung)
    cot_price_pay = ((int(cot_raw) + int(giam_tru_raw)) // 50000) * 50000
    vhkt_price_pay = (int(vhkt_raw) // 50000) * 50000 if vhkt_raw > 0 else 0
    
    # Lớp 2: Tính tổng tạm thời và làm tròn xuống cho TỔNG CHỐT
    temp_total = mb_raw + pm_raw + mfd_raw + cot_price_pay + vhkt_price_pay
    
    provided_new_price = int(site_data.get("new_price", 0))
    if provided_new_price > 0:
        new_price_val = provided_new_price
    else:
        # If pure extension, keep old price
        if template_type == "phu_luc_gia_han":
            new_price_val = float(site_data.get("old_price", 0) or 0)
        else:
            new_price_val = (int(temp_total) // 50000) * 50000
    
    # Lớp 3: Làm tròn phòng máy và dồn phần còn lại vào Mặt bằng
    pm_price_pay = (int(pm_raw) // 50000) * 50000 if pm_raw > 0 else 0
    mfd_price_pay = (int(mfd_raw) // 50000) * 50000 if mfd_raw > 0 else 0
    mb_price_pay = new_price_val - pm_price_pay - mfd_price_pay - cot_price_pay - vhkt_price_pay

    
    new_price_str = format_vn_currency(new_price_val)
    new_price_text = number_to_vietnamese_words(new_price_val) + " đồng"
    
    # Calculate Deduction Text
    deduction_text = ""
    total_deduction = 0
    paid_until = str(site_data.get("paid_until_date", "")).strip()
    if paid_until and paid_until not in ["nan", "None"]:
        try:
            paid_dt = datetime.strptime(paid_until[:10], "%d/%m/%Y")
            start_dt = datetime(2025, 10, 1)
            if paid_dt >= start_dt:
                # Calculate number of months overpaid
                # We add 1 day to paid_dt to make it the 1st of the next month (e.g. 31/10/2025 -> 01/11/2025)
                # so relativedelta returns exact months without fractions if it's a full month.
                paid_dt_adjusted = paid_dt + relativedelta(days=1)
                diff = relativedelta(paid_dt_adjusted, start_dt)
                months_overpaid = diff.years * 12 + diff.months + (diff.days / 30.0)
                
                # Calculate amount diff per month
                old_p = float(site_data.get("old_price", 0) or 0)
                new_p = float(new_price_val)
                diff_per_month = old_p - new_p
                
                if diff_per_month > 0 and months_overpaid > 0:
                    total_deduction = diff_per_month * months_overpaid
                    deduction_text = f"Giảm trừ số tiền {format_vn_currency(total_deduction)} đồng đã thanh toán từ 01/10/2025 đến {paid_until} vào kỳ thanh toán tiếp theo."
                else:
                    deduction_text = f"Giảm trừ số tiền đã thanh toán từ 01/10/2025 đến {paid_until} vào kỳ thanh toán tiếp theo."
        except Exception:
            deduction_text = f"Giảm trừ số tiền đã thanh toán từ 01/10/2025 đến {paid_until} vào kỳ thanh toán tiếp theo."
    
    # Generate CSHT List dynamically
    csht_list = (
        "- Vị trí đặt thiết bị: phòng có sẵn của chủ trạm, diện tích 3x3m, theo tiêu chuẩn của MobiFone.\n"
        "- Vị trí đặt máy nổ: phòng có sẵn của chủ trạm, diện tích 3mx2m, theo tiêu chuẩn của MobiFone.\n"
        "- Cột anten: cột dây co dưới mặt đất cao 36m có sẵn của chủ trạm.\n"
        "- Hệ thống tiếp đất và hệ thống lạnh: MobiFone tự trang bị."
    )

    # For pure extension, remove certificate basis
    hs_phap_ly = site_data.get("hs_phap_ly", "")
    if template_type == "phu_luc_gia_han":
        certificate_text = ""
    else:
        certificate_text = hs_phap_ly if hs_phap_ly and hs_phap_ly != "nan" else "................................................................................................................................"

    replacements = {
        "{{CERTIFICATE}}": certificate_text,
        "{{ CERTIFICATE }}": certificate_text,
        "{{SITE_ID}}": site_id,
        "{{ SITE_ID }}": site_id,
        "{{OWNER_NAME}}": owner_name,
        "{{ OWNER_NAME }}": owner_name,
        "{{ CONTRACT_NO }}": "",
        "{{CONTRACT_DATE}}": "",
        "{{ CONTRACT_DATE }}": "",
        "{{ADDRESS_OLD}}": address_old,
        "{{ ADDRESS_OLD }}": address_old,
        "{{ADDRESS_NEW}}": site_data.get("address_new", ""),
        "{{ ADDRESS_NEW }}": site_data.get("address_new", ""),
        " ({{ADDRESS_NEW}})": f" ({site_data.get('address_new')})" if site_data.get("address_new") and site_data.get("address_new") != "nan" else "",
        " ({{ ADDRESS_NEW }})": f" ({site_data.get('address_new')})" if site_data.get("address_new") and site_data.get("address_new") != "nan" else "",
        "{{LONGITUDE}}": longitude,
        "{{LATITUDE}}": latitude,
        "{{CERTIFICATE}}": certificate_text,
        "{{ CERTIFICATE }}": certificate_text,
        "{{LEGAL_DOCS}}": certificate_text,
        "{{ LEGAL_DOCS }}": certificate_text,
        "{{START_DATE}}": start_date_str,
        "{{END_DATE}}": end_date_str,
        "{{NEW_PRICE}}": new_price_str,
        "{{ NEW_PRICE }}": new_price_str,
        "{{NEW_PRICE_TEXT}}": new_price_text,
        "{{ NEW_PRICE_TEXT }}": new_price_text,
        "{{ACCOUNT_OWNER}}": account_owner,
        "{{ ACCOUNT_OWNER }}": account_owner,
        "{{ACCOUNT_NO}}": account_no,
        "{{ ACCOUNT_NO }}": account_no,
        "{{BANK_NAME}}": full_bank_name,
        "{{ BANK_NAME }}": full_bank_name,
        
        "{{MB_QĐ02}}": format_vn_currency(mb_raw),
        "{{ MB_QĐ02 }}": format_vn_currency(mb_raw),
        "{{P_MB}}": format_vn_currency(mb_price_pay),
        "{{ P_MB }}": format_vn_currency(mb_price_pay),
        "{{TL_MB}}": f"{((mb_price_pay / mb_raw) - 1) * 100:.2f}%" if mb_raw != 0 else "0%",
        "{{ TL_MB }}": f"{((mb_price_pay / mb_raw) - 1) * 100:.2f}%" if mb_raw != 0 else "0%",
        
        "{{PM_1245}}": format_vn_currency(pm_raw),
        "{{ PM_1245 }}": format_vn_currency(pm_raw),
        "{{P_PM}}": format_vn_currency(pm_price_pay),
        "{{ P_PM }}": format_vn_currency(pm_price_pay),
        "{{TL_PM}}": f"{((pm_price_pay / pm_raw) - 1) * 100:.2f}%" if pm_raw != 0 else "0%",
        "{{ TL_PM }}": f"{((pm_price_pay / pm_raw) - 1) * 100:.2f}%" if pm_raw != 0 else "0%",
        
        "{{MFĐ_1245}}": format_vn_currency(mfd_raw),
        "{{ MFĐ_1245 }}": format_vn_currency(mfd_raw),
        "{{P_MFD}}": format_vn_currency(mfd_price_pay),
        "{{ P_MFD }}": format_vn_currency(mfd_price_pay),
        "{{TL_MFD}}": f"{((mfd_price_pay / mfd_raw) - 1) * 100:.2f}%" if mfd_raw != 0 else "0%",
        "{{ TL_MFD }}": f"{((mfd_price_pay / mfd_raw) - 1) * 100:.2f}%" if mfd_raw != 0 else "0%",
        
        "{{COT_1245}}": format_vn_currency(cot_raw),
        "{{ COT_1245 }}": format_vn_currency(cot_raw),
        "{{GIAM_TRU}}": format_vn_currency(giam_tru_raw),
        "{{ GIAM_TRU }}": format_vn_currency(giam_tru_raw),
        "{{COT_CHOT}}": format_vn_currency(cot_price_pay),
        "{{ COT_CHOT }}": format_vn_currency(cot_price_pay),
        "{{TL_COT}}": f"{((cot_price_pay / (cot_raw + giam_tru_raw)) - 1) * 100:.2f}%" if (cot_raw + giam_tru_raw) != 0 else "0%",
        "{{ TL_COT }}": f"{((cot_price_pay / (cot_raw + giam_tru_raw)) - 1) * 100:.2f}%" if (cot_raw + giam_tru_raw) != 0 else "0%",
        
        "{{VHKT_1245}}": format_vn_currency(vhkt_raw),
        "{{ VHKT_1245 }}": format_vn_currency(vhkt_raw),
        "{{P_VHKT}}": format_vn_currency(vhkt_price_pay),
        "{{ P_VHKT }}": format_vn_currency(vhkt_price_pay),
        "{{TL_VHKT}}": f"{((vhkt_price_pay / vhkt_raw) - 1) * 100:.2f}%" if vhkt_raw != 0 else "0%",
        "{{ TL_VHKT }}": f"{((vhkt_price_pay / vhkt_raw) - 1) * 100:.2f}%" if vhkt_raw != 0 else "0%",
        
        "{{TONG_QD}}": format_vn_currency(mb_raw + pm_raw + mfd_raw + cot_raw + giam_tru_raw + vhkt_raw),
        "{{ TONG_QD }}": format_vn_currency(mb_raw + pm_raw + mfd_raw + cot_raw + giam_tru_raw + vhkt_raw),
        "{{TONG_CHOT}}": format_vn_currency(new_price_val),
        "{{ TONG_CHOT }}": format_vn_currency(new_price_val),
        "{{TL_TONG}}": f"{((new_price_val / (mb_raw + pm_raw + mfd_raw + cot_raw + giam_tru_raw + vhkt_raw)) - 1) * 100:.2f}%" if (mb_raw + pm_raw + mfd_raw + cot_raw + giam_tru_raw + vhkt_raw) != 0 else "0%",
        
        # Additional tags for Phu Luc Giam Gia and Thanh Ly
        "{{OLD_PRICE}}": format_vn_currency(site_data.get("old_price", 0)),
        "{{ OLD_PRICE }}": format_vn_currency(site_data.get("old_price", 0)),
        "{{PHONE}}": site_data.get("phone", ""),
        "{{ PHONE }}": site_data.get("phone", ""),
        "{{CONTACT_ADDR}}": site_data.get("contact_addr", ""),
        "{{ CONTACT_ADDR }}": site_data.get("contact_addr", ""),
        "{{CONTRACT_NO}}": site_data.get("contract_no", ""),
        "{{ CONTRACT_NO }}": site_data.get("contract_no", ""),
        "{{CONTRACT_DATE}}": site_data.get("contract_date", ""),
        "{{ CONTRACT_DATE }}": site_data.get("contract_date", ""),
        "{{PAID_UNTIL_DATE}}": paid_until,
        "{{ PAID_UNTIL_DATE }}": paid_until,
        "{{OLD_END_DATE}}": old_end_date_str,
        "{{ OLD_END_DATE }}": old_end_date_str,
        "{{CSHT_LIST1}}": csht_list,
        "{{ CSHT_LIST1 }}": csht_list,
        "{{CSHT_LIST2}}": csht_list,
        "{{ CSHT_LIST2 }}": csht_list,
        "{{DEDUCTION_TEXT}}": deduction_text,
        "{{ DEDUCTION_TEXT }}": deduction_text,
        "{{REPORT_DATE}}": datetime.now().strftime('%d/%m/%Y'),
        "{{ REPORT_DATE }}": datetime.now().strftime('%d/%m/%Y'),
        "{{TOTAL_AMOUNT}}": format_vn_currency(new_price_val * 6), # Will be overridden if needed
        "{{ TOTAL_AMOUNT }}": format_vn_currency(new_price_val * 6),
        "{{CYCLES_REMAINING}}": "0",
        "{{ CYCLES_REMAINING }}": "0",
        "{{OWNER_MATCH}}": banking.get("match_status_text", ""),
        "{{ OWNER_MATCH }}": banking.get("match_status_text", ""),
        "{{EXTENSION_STATUS}}": site_data.get("ext_status", ""),
        "{{ EXTENSION_STATUS }}": site_data.get("ext_status", "")
    }

    doc = docx.Document(template_path)
    
    # Remove DEDUCTION_TEXT paragraph if empty to avoid floating hyphen
    if not deduction_text:
        for p in doc.paragraphs:
            if "{{DEDUCTION_TEXT}}" in p.text or "{{ DEDUCTION_TEXT }}" in p.text:
                p_element = p._element
                p_element.getparent().remove(p_element)
                p._p = p._element = None
    
    # 5. GENERATE PAYMENT TABLE (Logic mới: Dồn ngày lẻ vào Kỳ 1)
    periods = []
    total_amount = 0
    cycle_months = 6
    term_years = 5

    # Lấy thông tin chu kỳ từ hàm tập trung
    # start_date_str đã được định nghĩa ở trên là start_date_dt.strftime("%d/%m/%Y")
    result = calculate_aligned_cycles(start_date_str, cycle_months, term_years, new_price_val)
    
    if "cycles" in result:
        for c in result["cycles"]:
            amt = c["amount"]
            if c["index"] == 1 and total_deduction > 0:
                amt -= total_deduction
                
            # Chuyển đổi ngược lại format để tương thích code cũ bên dưới
            periods.append({
                "no": c["index"],
                "start": datetime.strptime(c["start"], "%d/%m/%Y"),
                "end": datetime.strptime(c["end"], "%d/%m/%Y"),
                "amount": amt
            })
            total_amount += amt
        
        end_contract = datetime.strptime(result["end_date"], "%d/%m/%Y")
        replacements["{{END_DATE}}"] = result["end_date"]
        replacements["{{TOTAL_AMOUNT}}"] = format_vn_currency(total_amount)
        replacements["{{ TOTAL_AMOUNT }}"] = format_vn_currency(total_amount)

    # 6. Global Replacements
    total_replacement = f"Tổng cộng:             {format_vn_currency(total_amount)}"
    import re
    
    has_page_broken = False
    for p in doc.paragraphs:
        if "Tổng cộng:" in p.text:
            if re.search(r'[\d\.]+', p.text):
                p.text = total_replacement
                p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                    run.font.size = docx.shared.Pt(14)
                continue
                
        text = p.text
        original_text = text
        for tag, rep in replacements.items():
            if tag in text:
                text = text.replace(tag, str(rep))
                
        if text != original_text:
            p.text = text
            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                run.font.size = docx.shared.Pt(14)
                
        bold_triggers = [
            "Thời điểm bắt đầu tính tiền",
            "Giảm trừ số tiền",
            "Một bên là: Ông/Bà",
            "Thông tin tài khoản nhận thanh toán là",
            "Tên tài khoản:",
            "Số tài khoản:",
            "Tại ngân hàng:",
            "Địa chỉ trạm:"
        ]
        signature_triggers = [
            "ĐẠI DIỆN MOBIFONE",
            "ĐẠI DIỆN BÊN THUÊ",
            "ĐẠI DIỆN BÊN CHO THUÊ",
            "XÁC NHẬN CỦA ĐƠN VỊ"
        ]
        if not has_page_broken and any(trigger in p.text.upper() for trigger in signature_triggers):
            p.paragraph_format.page_break_before = True
            has_page_broken = True
            
        if any(trigger in p.text for trigger in bold_triggers) or (owner_name and p.text.strip() == owner_name.strip()):
            for run in p.runs:
                run.bold = True

    for table in doc.tables:
        for row_obj in table.rows:
            for cell in row_obj.cells:
                for p in cell.paragraphs:
                    text = p.text
                    original_text = text
                    for tag, rep in replacements.items():
                        if tag in text:
                            text = text.replace(tag, str(rep))
                            
                    if text != original_text:
                        p.text = text
                        for run in p.runs:
                            run.font.name = "Times New Roman"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                            run.font.size = docx.shared.Pt(14)
                            
                    if any(trigger in p.text for trigger in bold_triggers) or (owner_name and p.text.strip() == owner_name.strip()):
                        for run in p.runs:
                            run.bold = True

    # 7. Render Payment Table
    pay_table = None
    for table in doc.tables:
        if len(table.rows) > 0:
            first_cell_text = table.rows[0].cells[0].text
            row_texts = [cell.text for cell in table.rows[0].cells]
            if "kỳ" in first_cell_text.lower() or "thanh toán" in first_cell_text.lower() or any("{{PAY_ROW}}" in t or "{{ PAY_ROW }}" in t for t in row_texts):
                pay_table = table
                break
                
    if pay_table is None and len(doc.tables) > 8:
        pay_table = doc.tables[8]
        if is_phu_luc and len(doc.tables) > 4:
            pay_table = doc.tables[4] # Usually table 4 in Phu Luc template
            
    if pay_table:
        try:
            pay_table.style = 'Table Grid'
        except Exception:
            pass # Ignore if style doesn't exist
            
        for i, period in enumerate(periods):
            line = f"+ Kỳ {period['no']}: từ ngày {period['start'].strftime('%d/%m/%Y')} đến ngày {period['end'].strftime('%d/%m/%Y')}, Số tiền là: {format_vn_currency(period['amount'])} đồng"
            
            if i == 0 and len(pay_table.rows) > 0 and (pay_table.rows[0].cells[0].text.strip() == "" or "{{PAY_ROW}}" in pay_table.rows[0].cells[0].text or "{{ PAY_ROW }}" in pay_table.rows[0].cells[0].text):
                row_obj = pay_table.rows[0]
            else:
                row_obj = pay_table.add_row()
                
            row_obj.height = docx.shared.Inches(0.3)
            row_obj.height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.EXACTLY
            
            cell = row_obj.cells[0]
            cell.text = line
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                run.font.size = docx.shared.Pt(13)
                
        # Thêm dòng tổng cộng vào cuối bảng pay_table
        if len(periods) > 0:
            total_row = pay_table.add_row()
            total_row.height = docx.shared.Inches(0.3)
            total_row.height_rule = docx.enum.table.WD_ROW_HEIGHT_RULE.EXACTLY
            total_cell = total_row.cells[0]
            total_cell.text = f"Tổng cộng: {format_vn_currency(total_amount)} đồng"
            tp = total_cell.paragraphs[0]
            tp.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
            tp.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
            for run in tp.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                run.font.size = docx.shared.Pt(13)
                run.bold = True
            
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    output_filename = f"{prefix}_{site_id}.docx"
    output_path = os.path.join(output_dir, output_filename)
    doc.save(output_path)
    return True, output_filename
