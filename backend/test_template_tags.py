from docx import Document

doc = Document("../templates/MASTER_TEMPLATE_VFINAL.docx")
tags = set()
for para in doc.paragraphs:
    import re
    found = re.findall(r'\{\{(\w+)\}\}', para.text)
    for f in found:
        tags.add(f)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                found = re.findall(r'\{\{(\w+)\}\}', para.text)
                for f in found:
                    tags.add(f)
print("Tags found in template:")
for t in sorted(tags):
    print(f"  {{{{ {t} }}}}")
