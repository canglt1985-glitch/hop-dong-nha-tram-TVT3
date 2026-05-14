import docx
from docx.shared import Pt
import glob

def change_font_size(filepath, size=13):
    print(f"Processing {filepath}...")
    doc = docx.Document(filepath)
    
    # Update normal style
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(size)
    
    # Update all runs in paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            run.font.size = Pt(size)
            
    # Update all runs in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(size)
                        
    doc.save(filepath)
    print(f"Saved {filepath} with font size {size}pt.")

if __name__ == "__main__":
    templates = glob.glob("templates/*.docx")
    for t in templates:
        # Ignore temp files starting with ~
        if "~$" not in t:
            change_font_size(t)
