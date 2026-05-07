import os
import sys
import pandas as pd
import docx
from datetime import datetime
from dateutil.relativedelta import relativedelta
from docx.oxml.ns import qn

# Enforce UTF-8 console output for Vietnamese
try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

# Source files
EXCEL_PATH = r"D:\Chuyen doi so\soan HD\DATA HOP DONG.xlsx"
TEMPLATE_PATH = r"D:\Chuyen doi so\soan HD\soan-thao-phu-luc-hd\templates\THANH LY KY LAI_TEMPLATE.docx"
OUTPUT_DIR = r"D:\Chuyen doi so\soan HD\soan-thao-phu-luc-hd\output"

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# --- VIETNAMESE NUMBER TO WORDS CONVERTER ---
def format_vn_currency(amount):
    """Formats float/int into standard Vietnamese dot-separated string."""
    try:
        if pd.isna(amount) or amount is None:
            return "0"
        return "{:,.0f}".format(float(amount)).replace(",", ".")
    except Exception:
        return "0"

def number_to_vietnamese_words(number):
    """Converts a number to Vietnamese words precisely for contract values without external library dependencies."""
    try:
        number = int(number)
    except:
        return ""
        
    if number == 0:
        return "Không"
        
    units = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]
    
    def read_three_digits(n, is_first=False):
        hundreds = n // 100
        tens = (n % 100) // 10
        ones = n % 10
        
        res = []
        if hundreds > 0:
            res.append(units[hundreds] + " trăm")
        elif not is_first:
            res.append("không trăm")
            
        if tens > 0:
            if tens == 1:
                res.append("mười")
            else:
                res.append(units[tens] + " mươi")
        elif (hundreds > 0 or not is_first) and ones > 0:
            res.append("lẻ")
            
        if ones > 0:
            if ones == 1 and tens > 1:
                res.append("mốt")
            elif ones == 5 and tens > 0:
                res.append("lăm")
            else:
                res.append(units[ones])
                    
        return " ".join(res)
        
    groups = []
    temp = number
    while temp > 0:
        groups.append(temp % 1000)
        temp //= 1000
        
    group_units = ["", "nghìn", "triệu", "tỷ", "nghìn tỷ", "triệu tỷ"]
    words = []
    
    for idx, val in enumerate(groups):
        if val > 0:
            part = read_three_digits(val, idx == len(groups) - 1)
            unit = group_units[idx]
            if unit:
                words.append(part + " " + unit)
            else:
                words.append(part)
                
    words.reverse()
    text = " ".join(words)
    text = " ".join(text.split()).strip()
    return text[0].upper() + text[1:]

# --- DATE MATH CALCULATIONS ---
def calculate_contract_dates(expiry_date_val):
    """
    1. Parse raw date.
    2. Push to the last day of the expiring month.
    3. Start Date = Expiring Month End + 1 day.
    4. End Date = Start Date + 5 years - 1 day.
    """
    try:
        if isinstance(expiry_date_val, datetime):
            exp_date = expiry_date_val
        else:
            # Parse from standard pandas formats
            exp_date = pd.to_datetime(expiry_date_val)
            
        if pd.isna(exp_date):
            raise ValueError("Invalid date value")
            
        # Push to last day of that month
        month_end_date = exp_date + relativedelta(day=31) # dateutil pushes to last valid day of month
        
        # New contract start date
        start_date = month_end_date + relativedelta(days=1)
        
        # New contract end date
        end_date = start_date + relativedelta(years=5) - relativedelta(days=1)
        
        return start_date, end_date
    except Exception as e:
        print(f"  [Date Warning] Using fallback dating: {e}")
        # Default fallback
        return datetime(2027, 7, 1), datetime(2032, 6, 30)

# --- RUN-PRESERVING TAG SUBSTITUTION ---
def replace_tag_in_runs(p, tag, replacement):
    """
    Replaces a specific tag in a paragraph's runs while keeping all bold, italic,
    underlines, and explicitly setting Font to Times New Roman.
    """
    if tag not in p.text:
        return False
        
    replaced = False
    
    # 1. Easy path: target is fully in one run
    for run in p.runs:
        if tag in run.text:
            run.text = run.text.replace(tag, str(replacement))
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
            replaced = True
            
    # 2. Split-run fallback
    if not replaced:
        # Merge runs carefully if tag is split (unlikely but possible)
        p.text = p.text.replace(tag, str(replacement))
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
        replaced = True
        
    return replaced

# --- RUN CONTRACT GENERATION ---
def generate_thanh_ly_contract(site_id="DNXL10"):
    print(f"🚀 Initializing Standardized contract generation for: {site_id}")
    
    # 1. Load Master Excel Sheet
    print(f"Loading Master Excel from: {EXCEL_PATH}")
    df = pd.read_excel(EXCEL_PATH, sheet_name="FULL-(XHH-LK)", header=None)
    
    # Find the target row index
    row_idx = None
    for r in range(df.shape[0]):
        if str(df.iloc[r, 1]).strip().upper() == site_id.upper():
            row_idx = r
            break
            
    if row_idx is None:
        print(f"❌ Error: Site {site_id} not found in Column B!")
        return False
        
    print(f"Found {site_id} at row: {row_idx + 1}")
    row = df.iloc[row_idx]
    
    # 2. Extract Data dynamically with safeguards
    owner_name = str(row.iloc[16]).strip() if pd.notna(row.iloc[16]) else ""
    contact_addr = str(row.iloc[17]).strip() if pd.notna(row.iloc[17]) else ""
    phone = str(row.iloc[18]).strip() if pd.notna(row.iloc[18]) else ""
    cccd = str(row.iloc[19]).strip() if pd.notna(row.iloc[19]) else ""
    
    contract_no = str(row.iloc[24]).strip() if pd.notna(row.iloc[24]) else ""
    raw_contract_date = row.iloc[25]
    if isinstance(raw_contract_date, datetime):
        contract_date = raw_contract_date.strftime("%d/%m/%Y")
    else:
        contract_date = str(raw_contract_date).strip() if pd.notna(raw_contract_date) else ""
        
    raw_expiry_date = row.iloc[26]
    start_date_dt, end_date_dt = calculate_contract_dates(raw_expiry_date)
    start_date_str = start_date_dt.strftime("%d/%m/%Y")
    end_date_str = end_date_dt.strftime("%d/%m/%Y")
    
    address_old = str(row.iloc[35]).strip() if pd.notna(row.iloc[35]) else ""
    
    # Safe check for Column 42 (Index 41): New Ward/Commune
    # In legacy sheet, Column 42 (Index 41) holds Ward. If the user added HPSL, let's detect it.
    col42_val = str(row.iloc[41]).strip() if pd.notna(row.iloc[41]) else ""
    
    # Safe ward mapping
    if "xã" in col42_val.lower() or "phường" in col42_val.lower() or "tỉnh" in col42_val.lower():
        # This is a ward, not a legal certificate! Output a blank line for manual write-in.
        certificate = "________________________________________________________________________________"
        address_new = col42_val
    else:
        # User has populated HSPL certificate! Use it directly.
        certificate = col42_val if col42_val else "________________________________________________________________________________"
        address_new = "Xã Tân An, Huyện Vĩnh Cửu, Tỉnh Đồng Nai" # Fallback new address
        
    # Coordinates
    longitude = str(row.iloc[96]).strip() if pd.notna(row.iloc[96]) else ""
    latitude = str(row.iloc[97]).strip() if pd.notna(row.iloc[97]) else ""
    
    # --- HARMONIZED PRICING EXTRACTION (Indices 47 to 67) ---
    mb_raw = float(row.iloc[47]) if pd.notna(row.iloc[47]) else 0.0
    
    # Room/Shelter types dynamic scanning (Indices 48 to 54)
    pm_raw = 0.0
    pm_label_val = "Phòng máy mặt đất"
    for col_i in range(48, 55):
        val = row.iloc[col_i]
        if pd.notna(val) and isinstance(val, (int, float)) and val > 0:
            pm_raw = float(val)
            try:
                title = str(df.iloc[6, col_i]).strip()
                if title and title.lower() != 'nan':
                    pm_label_val = title.capitalize()
            except Exception:
                pass
            break
            
    # Generator Room (Index 55)
    mfd_raw = float(row.iloc[55]) if pd.notna(row.iloc[55]) else 0.0
    
    # Tower types dynamic scanning (Indices 56 to 58)
    cot_raw = 0.0
    for col_i in range(56, 59):
        val = row.iloc[col_i]
        if pd.notna(val) and isinstance(val, (int, float)) and val > 0:
            cot_raw = float(val)
            break
            
    # Shared Tower Discount (Index 67)
    giam_tru_raw = float(row.iloc[67]) if pd.notna(row.iloc[67]) else 0.0
    
    # Auxiliary Extra Items scanning (Indices 59 to 66)
    other_items = []
    other_items_raw_total = 0.0
    other_items_pay_total = 0.0
    for col_i in range(59, 67):
        val = row.iloc[col_i]
        if pd.notna(val) and isinstance(val, (int, float)) and val > 0:
            raw_val = float(val)
            pay_val = (int(raw_val) // 50000) * 50000
            try:
                title = str(df.iloc[6, col_i]).strip()
            except Exception:
                title = ""
            if not title or title.lower() == 'nan':
                title = f"Hạng mục Index {col_i}"
            other_items.append({
                "title": title.capitalize(),
                "raw": raw_val,
                "pay": pay_val
            })
            other_items_raw_total += raw_val
            other_items_pay_total += pay_val
            
    # --- 3-LAYER ROUNDING MATHS (Unified with VB1245) ---
    # Layer 1: Round Antenna Tower price including negative sharing discounts
    cot_price_pay = ((int(cot_raw) + int(giam_tru_raw)) // 50000) * 50000
    
    # Layer 2: Temporary sum rounded down to nearest 50,000 for standard compliance
    temp_total = mb_raw + pm_raw + mfd_raw + cot_price_pay + other_items_raw_total
    new_price_val = (int(temp_total) // 50000) * 50000
    
    # Layer 3: Round room/generator prices and plug the remaining portion into Mặt bằng
    pm_price_pay = (int(pm_raw) // 50000) * 50000 if pm_raw > 0 else 0
    mfd_price_pay = (int(mfd_raw) // 50000) * 50000 if mfd_raw > 0 else 0
    mb_price_pay = new_price_val - pm_price_pay - mfd_price_pay - cot_price_pay - other_items_pay_total
    
    new_price_str = format_vn_currency(new_price_val)
    new_price_text = number_to_vietnamese_words(new_price_val) + " đồng"
    
    # Bank Account details
    account_owner = str(row.iloc[87]).strip() if pd.notna(row.iloc[87]) else owner_name
    account_no = str(row.iloc[88]).strip() if pd.notna(row.iloc[88]) else ""
    bank_name = str(row.iloc[89]).strip() if pd.notna(row.iloc[89]) else ""
    bank_branch = str(row.iloc[90]).strip() if pd.notna(row.iloc[90]) else ""
    full_bank_name = f"{bank_name} - {bank_branch}" if bank_branch else bank_name
    
    print("--- EXTRACTED & ROUNDED DETAILS ---")
    print(f"  Owner: {owner_name} | Phone: {phone} | CCCD: {cccd}")
    print(f"  Contract: {contract_no} Signed: {contract_date} Expiry: {raw_expiry_date}")
    print(f"  Calculated Start: {start_date_str} Calculated End: {end_date_str}")
    print(f"  Old Address: {address_old}")
    print(f"  New Address: {address_new}")
    print(f"  Coordinates: Long {longitude} - Lat {latitude}")
    print(f"  HPSL (Certificate): {certificate[:50]}...")
    print(f"  PM Label: {pm_label_val} | PM Raw: {pm_raw} -> Round: {pm_price_pay}")
    print(f"  MB Raw: {mb_raw} -> Adjusted pay: {mb_price_pay}")
    print(f"  Tower Raw: {cot_raw} | Shared Disc: {giam_tru_raw} -> Round: {cot_price_pay}")
    print(f"  Other items raw total: {other_items_raw_total} -> Round: {other_items_pay_total}")
    print(f"  Final Price: {new_price_str} VND ({new_price_text})")
    print(f"  Bank: {account_owner} | A/C: {account_no} | Bank: {full_bank_name}")
    print("-------------------------")
    
    # 3. Create Replacement Dict
    replacements = {
        "{{SITE_ID}}": site_id,
        "{{OWNER_NAME}}": owner_name,
        "{{CONTACT_ADDR}}": contact_addr,
        "{{PHONE}}": phone,
        "{{CCCD}}": cccd,
        "{{CONTRACT_NO}}": contract_no,
        "{{CONTRACT_DATE}}": contract_date,
        "{{ADDRESS_OLD}}": address_old,
        "{{ADDRESS_NEW}}": address_new,
        "{{LONGITUDE}}": longitude,
        "{{LATITUDE}}": latitude,
        "{{CERTIFICATE}}": certificate,
        "{{START_DATE}}": start_date_str,
        "{{END_DATE}}": end_date_str,
        "{{NEW_PRICE}}": new_price_str,
        "{{NEW_PRICE_TEXT}}": new_price_text,
        "{{ACCOUNT_OWNER}}": account_owner,
        "{{ACCOUNT_NO}}": account_no,
        "{{BANK_NAME}}": full_bank_name,
        
        # Harmonized Excel table cell tags
        "{{MB_QĐ02}}": format_vn_currency(mb_raw),
        "{{P_MB}}": format_vn_currency(mb_price_pay),
        "{{TL_MB}}": f"{((mb_price_pay / mb_raw) - 1) * 100:.2f}%" if mb_raw != 0 else "0%",
        
        "{{PM_1245}}": format_vn_currency(pm_raw),
        "{{P_PM}}": format_vn_currency(pm_price_pay),
        "{{TL_PM}}": f"{((pm_price_pay / pm_raw) - 1) * 100:.2f}%" if pm_raw != 0 else "0%",
        "{{PM_LABEL}}": pm_label_val,
        "Phòng máy mặt đất": pm_label_val,
        
        "{{MFĐ_1245}}": format_vn_currency(mfd_raw),
        "{{P_MFD}}": format_vn_currency(mfd_price_pay),
        "{{TL_MFD}}": f"{((mfd_price_pay / mfd_raw) - 1) * 100:.2f}%" if mfd_raw != 0 else "0%",
        
        "{{COT_1245}}": format_vn_currency(cot_raw),
        "{{GIAM_TRU}}": format_vn_currency(giam_tru_raw),
        "{{COT_CHOT}}": format_vn_currency(cot_price_pay),
        "{{TL_COT}}": f"{((cot_price_pay / (cot_raw + giam_tru_raw)) - 1) * 100:.2f}%" if (cot_raw + giam_tru_raw) != 0 else "0%",
        
        "{{TONG_QD}}": format_vn_currency(mb_raw + pm_raw + mfd_raw + cot_raw + giam_tru_raw + other_items_raw_total),
        "{{TONG_CHOT}}": format_vn_currency(new_price_val),
        "{{TL_TONG}}": f"{((new_price_val / (mb_raw + pm_raw + mfd_raw + cot_raw + giam_tru_raw + other_items_raw_total)) - 1) * 100:.2f}%" if (mb_raw + pm_raw + mfd_raw + cot_raw + giam_tru_raw + other_items_raw_total) != 0 else "0%"
    }
    
    # 4. Load Word Template
    doc = docx.Document(TEMPLATE_PATH)
    
    # 5. Dynamically Append Extra Items into Price Table (Table 1)
    if other_items:
        price_table = doc.tables[1]
        for item in other_items:
            # Add row and insert before 'Tổng cộng' row (last row of table 1)
            new_row = price_table.add_row()
            price_table._tbl.insert(-2, new_row._tr)
            
            # Fill cells
            new_row.cells[0].text = item["title"]
            new_row.cells[1].text = ""
            new_row.cells[2].text = format_vn_currency(item["raw"])
            new_row.cells[3].text = ""
            new_row.cells[4].text = format_vn_currency(item["pay"])
            ratio_val = f"{((item['pay'] / item['raw']) - 1) * 100:.2f}%" if item['raw'] != 0 else "0%"
            new_row.cells[5].text = ratio_val
            
            # Format text formatting inside new row cells to Times New Roman 14
            for cell in new_row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
                    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                        run.font.size = docx.shared.Pt(14)
                        
    # 6. Perform Run-Preserving Replacement
    for p in doc.paragraphs:
        for tag, rep in replacements.items():
            replace_tag_in_runs(p, tag, rep)
            
    # Perform replacement in all tables as well
    for table in doc.tables:
        for row_obj in table.rows:
            for cell in row_obj.cells:
                for p in cell.paragraphs:
                    for tag, rep in replacements.items():
                        replace_tag_in_runs(p, tag, rep)
                        
    # 7. Build the 10-period, 6-month payment cycle table dynamically!
    pay_table = None
    for table in doc.tables:
        if len(table.rows) > 0:
            first_cell_text = table.rows[0].cells[0].text
            row_texts = [cell.text for cell in table.rows[0].cells]
            if "kỳ" in first_cell_text.lower() or "thanh toán" in first_cell_text.lower() or any("{{PAY_ROW}}" in t for t in row_texts):
                pay_table = table
                break
                
    if pay_table is None:
        if len(doc.tables) > 8:
            pay_table = doc.tables[8] # The newly scanned index is table 8
            
    if pay_table:
        print("Located payment schedule table. Generating 10 periods...")
        curr_p_start = start_date_dt
        cycle_months = 6
        cycle_amount = new_price_val * cycle_months
        cycle_amount_str = format_vn_currency(cycle_amount)
        
        for i in range(10):
            curr_p_end = curr_p_start + relativedelta(months=cycle_months) - relativedelta(days=1)
            line = f"+ Kỳ {i+1}: từ ngày {curr_p_start.strftime('%d/%m/%Y')} đến ngày {curr_p_end.strftime('%d/%m/%Y')} Số tiền là: {cycle_amount_str} đồng"
            
            if i == 0 and len(pay_table.rows) > 0 and (pay_table.rows[0].cells[0].text.strip() == "" or "{{PAY_ROW}}" in pay_table.rows[0].cells[0].text):
                row_obj = pay_table.rows[0]
            else:
                row_obj = pay_table.add_row()
                
            cell = row_obj.cells[0]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            
            run = p.add_run(line)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
            run.font.size = docx.shared.Pt(14)
            
            # Next cycle start
            curr_p_start = curr_p_end + relativedelta(days=1)
            
    # 8. Save output
    output_filename = f"Thanh_Ly_Ky_Lai_{site_id}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"🎉 SUCCESS! Completed compliant contract generated at: {output_path}")
    return True

if __name__ == "__main__":
    generate_thanh_ly_contract("DNXL10")
