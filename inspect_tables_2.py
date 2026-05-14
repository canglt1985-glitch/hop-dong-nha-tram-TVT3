import docx

doc_c = docx.Document("templates/PHU_LUC_GIAM_GIA_CSHT.docx")
print("=== CSHT Table 1 ===")
for r in doc_c.tables[1].rows:
    print(" | ".join([c.text.strip().replace("\n", " ") for c in r.cells]))

doc_m = docx.Document("templates/PHU_LUC_GIAM_GIA_MAT_BANG.docx")
print("\n=== MAT_BANG Table 1 ===")
for r in doc_m.tables[1].rows:
    print(" | ".join([c.text.strip().replace("\n", " ") for c in r.cells]))

