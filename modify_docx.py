import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

def insert_row_above(table, row_idx):
    reference_row = table.rows[row_idx]
    new_row = table.add_row()
    # Move the new row to above the reference_row
    tbl = table._tbl
    tbl.insert(row_idx, new_row._tr)
    # Copy styling from the reference row
    for i, cell in enumerate(new_row.cells):
        for run in reference_row.cells[i].paragraphs[0].runs:
            cell.paragraphs[0].paragraph_format.alignment = reference_row.cells[i].paragraphs[0].paragraph_format.alignment
            new_run = cell.paragraphs[0].add_run()
            new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.bold:
                new_run.font.bold = run.font.bold
    return new_row

def update_template(filename):
    doc = docx.Document(filename)
    table = doc.tables[1]
    
    # Check if we already inserted it to avoid duplicates
    for row in table.rows:
        if "Bảo vệ, hỗ trợ VHKT, PCCC" in row.cells[0].text:
            print(f"Already updated {filename}")
            return
            
    # Insert before 'Tổng cộng' which is usually the last row
    idx = len(table.rows) - 1
    new_row = insert_row_above(table, idx)
    
    values = ['Bảo vệ, hỗ trợ VHKT, PCCC', '', '{{VHKT_1245}}', '', '{{P_VHKT}}', '{{TL_VHKT}}']
    for i, cell in enumerate(new_row.cells):
        cell.text = values[i]
        # Keep font Times New Roman
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                
    doc.save(filename)
    print(f"Updated {filename}")

update_template("templates/PHU_LUC_GIAM_GIA_MAT_BANG.docx")
update_template("templates/THANH_LY_KY_LAI_MAT_BANG.docx")
