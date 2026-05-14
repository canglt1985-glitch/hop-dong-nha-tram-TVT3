import docx
import copy
from docx.oxml.ns import qn

def insert_row(table, insert_idx, texts):
    tbl = table._tbl
    template_tr = table.rows[1]._tr # use row 1 as template (e.g. Mặt bằng)
    new_tr = copy.deepcopy(template_tr)
    tbl.insert(insert_idx, new_tr)
    
    # After inserting, python-docx automatically recognizes it if we re-fetch table.rows
    # but let's be careful
    new_row = table.rows[insert_idx]
    for idx, text in enumerate(texts):
        new_row.cells[idx].text = text
        for p in new_row.cells[idx].paragraphs:
            for run in p.runs:
                run.font.name = "Times New Roman"

def replace_30_09_2028(doc):
    for p in doc.paragraphs:
        if "30/09/2028" in p.text:
            p.text = p.text.replace("30/09/2028", "{{END_DATE}}")
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if "30/09/2028" in p.text:
                        p.text = p.text.replace("30/09/2028", "{{END_DATE}}")

def insert_deduction_text(doc):
    # Find {{PAY_ROW}}
    for idx, p in enumerate(doc.paragraphs):
        if "{{PAY_ROW}}" in p.text or "Kỳ 1: từ ngày" in p.text:
            p.insert_paragraph_before("{{DEDUCTION_TEXT}}")
            return
    for table in doc.tables:
        if len(table.rows) > 0 and "{{PAY_ROW}}" in table.rows[0].cells[0].text:
            # Add a paragraph before the table
            p = table._element.getprevious()
            # It's an oxml element, not easy. 
            pass

def patch():
    files = [
        "templates/PHU_LUC_GIAM_GIA_CSHT.docx",
        "templates/PHU_LUC_GIAM_GIA_MAT_BANG.docx",
        "templates/THANH_LY_KY_LAI_CSHT.docx",
        "templates/THANH_LY_KY_LAI_MAT_BANG.docx"
    ]
    
    for f in files:
        doc = docx.Document(f)
        replace_30_09_2028(doc)
        
        # Check if DEDUCTION_TEXT exists
        has_deduction = any("{{DEDUCTION_TEXT}}" in p.text for p in doc.paragraphs)
        if not has_deduction and "THANH_LY" in f:
            # We add DEDUCTION_TEXT right at the end of "Điều 3: Thanh toán"
            # It's usually before the payment table
            found = False
            for p in doc.paragraphs:
                if "tiền thuê trạm sẽ được tính từ ngày" in p.text or "Thời điểm bắt đầu tính tiền thuê" in p.text:
                    new_p = p.insert_paragraph_before("{{DEDUCTION_TEXT}}")
                    new_p.style = p.style
                    found = True
                    break
            if not found:
                print(f"Could not find insert point for DEDUCTION_TEXT in {f}")

        # Table 1 patch
        tbl = doc.tables[1]
        texts_in_col0 = [r.cells[0].text.strip() for r in tbl.rows]
        
        if "CSHT" in f and "Bảo vệ, hỗ trợ VHKT, PCCC" not in texts_in_col0:
            insert_idx = len(tbl.rows) - 1 # before "Tổng cộng"
            insert_row(tbl, insert_idx, ["Bảo vệ, hỗ trợ VHKT, PCCC", "", "{{VHKT_1245}}", "", "{{P_VHKT}}", "{{TL_VHKT}}"])
            
        if "MAT_BANG" in f and "Phòng MFĐ" not in texts_in_col0:
            insert_idx = 2 # after "Mặt bằng"
            insert_row(tbl, insert_idx, ["Phòng MFĐ", "", "{{MFĐ_1245}}", "", "{{P_MFD}}", "{{TL_MFD}}"])
            
        doc.save(f)
        print(f"Patched {f}")

patch()
